import os
import shutil
import json
from typing import Optional, List, Dict, Any
from fastapi import UploadFile, HTTPException
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_core.documents import Document

try:
    from langchain_classic.chains import RetrievalQA
except ImportError:
    from langchain.chains import RetrievalQA

from core.config import settings
from core.services.model_service import get_embeddings, get_llm
from core.schemas import Citation, QueryResponse

class RAGService:
    def __init__(self):
        self.vector_db: Optional[Chroma] = None
        self.docs_metadata_path = os.path.join(settings.DATA_DIR, "indexed_docs.json")
        self.indexed_documents: List[Dict[str, Any]] = self._load_indexed_documents()

    def _load_indexed_documents(self) -> List[Dict[str, Any]]:
        if os.path.exists(self.docs_metadata_path):
            try:
                with open(self.docs_metadata_path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    if isinstance(data, list):
                        return data
            except Exception as e:
                print(f"[Warning] Failed to load indexed_docs.json: {e}")
        return []

    def _save_indexed_documents(self):
        try:
            os.makedirs(os.path.dirname(self.docs_metadata_path), exist_ok=True)
            with open(self.docs_metadata_path, "w", encoding="utf-8") as f:
                json.dump(self.indexed_documents, f, indent=2)
        except Exception as e:
            print(f"[Warning] Failed to save indexed_docs.json: {e}")

    def _load_file_documents(self, file_path: str, ext: str, safe_filename: str) -> List[Document]:
        documents: List[Document] = []
        if ext == '.pdf':
            try:
                loader = PyPDFLoader(file_path)
                documents = loader.load()
            except Exception:
                try:
                    import pypdf
                    reader = pypdf.PdfReader(file_path)
                    for i, page in enumerate(reader.pages):
                        txt = page.extract_text() or ""
                        if txt.strip():
                            documents.append(Document(page_content=txt, metadata={"source": safe_filename, "page": i}))
                except Exception:
                    pass
        elif ext == '.docx':
            try:
                import docx
                doc = docx.Document(file_path)
                lines = []
                for p in doc.paragraphs:
                    if p.text.strip():
                        lines.append(p.text.strip())
                for table in doc.tables:
                    for row in table.rows:
                        row_txt = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                        if row_txt:
                            lines.append(row_txt)
                combined_text = "\n".join(lines)
                if combined_text.strip():
                    documents = [Document(page_content=combined_text, metadata={"source": safe_filename, "page": 0})]
            except Exception:
                try:
                    loader = TextLoader(file_path, encoding='utf-8')
                    documents = loader.load()
                except Exception:
                    pass
        elif ext in ['.txt', '.md', '.json', '.html', '.rtf', '.csv']:
            try:
                loader = TextLoader(file_path, encoding='utf-8')
                documents = loader.load()
            except Exception:
                try:
                    loader = TextLoader(file_path, encoding='latin-1')
                    documents = loader.load()
                except Exception:
                    pass
        return documents

    def get_or_create_vector_db(self) -> Chroma:
        if self.vector_db is None:
            self.vector_db = Chroma(
                persist_directory=settings.CHROMA_DIR,
                embedding_function=get_embeddings()
            )
        return self.vector_db

    async def process_file_upload(self, file: UploadFile) -> Dict[str, Any]:
        raw_filename = file.filename or "uploaded_document.pdf"
        safe_filename = os.path.basename(raw_filename)
        ext = os.path.splitext(safe_filename)[1].lower()

        allowed_extensions = ['.pdf', '.txt', '.md', '.docx', '.csv', '.json', '.html', '.rtf']
        if ext not in allowed_extensions:
            raise HTTPException(
                status_code=400, 
                detail=f"Format '{ext}' is not supported. Supported formats: .pdf, .docx, .txt, .md, .csv, .json, .html, .rtf"
            )

        file_path = os.path.abspath(os.path.join(settings.UPLOADS_DIR, safe_filename))
        if not file_path.startswith(os.path.abspath(settings.UPLOADS_DIR)):
            raise HTTPException(status_code=403, detail="Security Violation: Invalid file path traversal detected.")

        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        file_size = os.path.getsize(file_path)
        if file_size > settings.MAX_FILE_SIZE_BYTES:
            os.remove(file_path)
            raise HTTPException(status_code=400, detail=f"File exceeds maximum size limit of {settings.MAX_FILE_SIZE_BYTES // (1024*1024)} MB.")

        try:
            documents = self._load_file_documents(file_path, ext, safe_filename)

            if not documents or not any(doc.page_content.strip() for doc in documents):
                raise HTTPException(
                    status_code=400, 
                    detail=f"Could not extract readable text from '{safe_filename}'. If it is a scanned PDF image, please use a searchable text document."
                )

            page_count = len(documents)
            splitter = RecursiveCharacterTextSplitter(
                chunk_size=500,
                chunk_overlap=50
            )
            chunks = splitter.split_documents(documents)

            db = self.get_or_create_vector_db()
            db.add_documents(chunks)

            doc_info = {
                "filename": safe_filename,
                "pages": page_count,
                "chunks": len(chunks),
                "file_size_bytes": file_size,
                "extension": ext
            }

            if not any(d["filename"] == safe_filename for d in self.indexed_documents):
                self.indexed_documents.append(doc_info)
            self._save_indexed_documents()

            return {
                "status": "success",
                "message": f"Successfully processed '{safe_filename}'",
                "document": doc_info,
                "total_indexed_documents": len(self.indexed_documents)
            }
        except HTTPException:
            if os.path.exists(file_path):
                os.remove(file_path)
            raise
        except Exception as e:
            if os.path.exists(file_path):
                os.remove(file_path)
            raise HTTPException(status_code=500, detail=f"Error processing '{safe_filename}': {str(e)}")

    async def process_raw_text_upload(self, title: str, content: str) -> Dict[str, Any]:
        sanitized_title = os.path.basename(title.strip() if title else "Pasted_Note")
        if not sanitized_title.endswith('.txt'):
            sanitized_title += ".txt"

        file_path = os.path.abspath(os.path.join(settings.UPLOADS_DIR, sanitized_title))
        if not file_path.startswith(os.path.abspath(settings.UPLOADS_DIR)):
            raise HTTPException(status_code=403, detail="Security Violation: Invalid file path traversal detected.")

        with open(file_path, "w", encoding="utf-8") as f:
            f.write(content)

        file_size = os.path.getsize(file_path)
        try:
            documents = [Document(page_content=content, metadata={"source": sanitized_title, "page": 0})]

            splitter = RecursiveCharacterTextSplitter(
                chunk_size=500,
                chunk_overlap=50
            )
            chunks = splitter.split_documents(documents)

            db = self.get_or_create_vector_db()
            db.add_documents(chunks)

            doc_info = {
                "filename": sanitized_title,
                "pages": 1,
                "chunks": len(chunks),
                "file_size_bytes": file_size,
                "extension": ".txt"
            }

            if not any(d["filename"] == sanitized_title for d in self.indexed_documents):
                self.indexed_documents.append(doc_info)
            self._save_indexed_documents()

            return {
                "status": "success",
                "message": f"Successfully added raw text note '{sanitized_title}'",
                "document": doc_info,
                "total_indexed_documents": len(self.indexed_documents)
            }
        except Exception as e:
            if os.path.exists(file_path):
                os.remove(file_path)
            raise HTTPException(status_code=500, detail=f"Error processing text note safely: {str(e)}")

    async def delete_single_document(self, filename: str) -> Dict[str, Any]:
        safe_filename = os.path.basename(filename)
        self.indexed_documents = [d for d in self.indexed_documents if d["filename"] != safe_filename]

        file_path = os.path.abspath(os.path.join(settings.UPLOADS_DIR, safe_filename))
        if os.path.exists(file_path) and file_path.startswith(os.path.abspath(settings.UPLOADS_DIR)):
            try:
                os.remove(file_path)
            except Exception as e:
                print(f"[Warning] Could not delete file {file_path}: {e}")

        if self.vector_db is not None:
            self.vector_db.delete_collection()
            self.vector_db = None
            self.get_or_create_vector_db()

        for doc_item in list(self.indexed_documents):
            remaining_file = os.path.join(settings.UPLOADS_DIR, doc_item["filename"])
            if os.path.exists(remaining_file):
                ext = doc_item.get("extension", os.path.splitext(doc_item["filename"])[1].lower())
                docs = self._load_file_documents(remaining_file, ext, doc_item["filename"])
                if docs:
                    chunks = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50).split_documents(docs)
                    self.vector_db.add_documents(chunks)

        self._save_indexed_documents()

        return {
            "status": "success",
            "message": f"Document '{safe_filename}' deleted successfully.",
            "remaining_documents": len(self.indexed_documents)
        }

    async def execute_query(
        self,
        question: str,
        chat_history: Optional[List[Dict[str, Any]]] = None,
        top_k: int = 3,
        max_length: int = 512
    ) -> QueryResponse:
        sanitized_question = question.strip() if question else ""
        if not sanitized_question:
            raise HTTPException(status_code=400, detail="Question cannot be empty.")

        if len(self.indexed_documents) == 0:
            return QueryResponse(
                question=sanitized_question,
                answer="No documents are currently indexed in your library. Please upload a PDF, Word document, or paste a note in the sidebar first.",
                citations=[]
            )

        try:
            db = self.get_or_create_vector_db()
            llm = get_llm()

            k_val = max(1, min(top_k, 10))
            retriever = db.as_retriever(search_kwargs={"k": k_val})

            history_context = ""
            if chat_history and len(chat_history) > 0:
                recent_turns = chat_history[-4:]
                formatted_turns = []
                for turn in recent_turns:
                    role = "User" if turn.get("sender") == "user" else "Assistant"
                    formatted_turns.append(f"{role}: {turn.get('text', '')}")
                history_context = "Previous Conversation Context:\n" + "\n".join(formatted_turns) + "\n\nCurrent Question: "

            augmented_query = f"{history_context}{sanitized_question}"
            retrieved_docs = retriever.invoke(sanitized_question)

            if not retrieved_docs:
                return QueryResponse(
                    question=sanitized_question,
                    answer="No relevant content found in the uploaded documents. Please upload a PDF or document first.",
                    citations=[]
                )

            qa_chain = RetrievalQA.from_chain_type(
                llm=llm,
                retriever=retriever,
                return_source_documents=True
            )

            result = qa_chain.invoke({"query": augmented_query})
            answer = result.get("result", "Unable to generate an answer.").strip()

            citations = []
            source_docs = result.get("source_documents", retrieved_docs)
            for doc in source_docs:
                meta = doc.metadata or {}
                source_name = os.path.basename(meta.get("source", "Uploaded Document"))
                page_num = int(meta.get("page", 0)) + 1
                content_snippet = doc.page_content.strip()
                citations.append(Citation(
                    source=source_name,
                    page=page_num,
                    content=content_snippet
                ))

            return QueryResponse(
                question=sanitized_question,
                answer=answer,
                citations=citations
            )
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error executing RAG query safely: {str(e)}")

    async def clear_all(self) -> Dict[str, Any]:
        try:
            if self.vector_db is not None:
                self.vector_db.delete_collection()
                self.vector_db = None
            
            self.indexed_documents.clear()
            self._save_indexed_documents()

            if os.path.exists(settings.UPLOADS_DIR):
                for fname in os.listdir(settings.UPLOADS_DIR):
                    fpath = os.path.join(settings.UPLOADS_DIR, fname)
                    if os.path.isfile(fpath):
                        os.remove(fpath)

            self.get_or_create_vector_db()

            return {
                "status": "success",
                "message": "Database and temporary storage cleared safely."
            }
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to clear database safely: {str(e)}")

rag_service = RAGService()
